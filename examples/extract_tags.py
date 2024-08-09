import re
from docx import Document

def extract_jinja_tags(docx_path):
    # Загружаем документ
    doc = Document(docx_path)
    
    # Регулярное выражение для поиска Jinja тегов
    jinja_pattern = r'\{\%.*?\%\}|\{\{.*?\}\}'
    
    jinja_tags = []
    
    # Обходим все параграфы в документе
    for para in doc.paragraphs:
        jinja_tags.extend(re.findall(jinja_pattern, para.text))
    
    # Обходим все таблицы в документе
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                jinja_tags.extend(re.findall(jinja_pattern, cell.text))

    # Обходим колонтитулы
    for section in doc.sections:
        for header in section.header.paragraphs:
            jinja_tags.extend(re.findall(jinja_pattern, header.text))
        for footer in section.footer.paragraphs:
            jinja_tags.extend(re.findall(jinja_pattern, footer.text))

    # Поиск в фигурах (шапках с рисованием, если поддерживаются)
    for shape in doc.inline_shapes:
        if shape.type == 3:  # 3 - тип 'Picture'
            continue
        jinja_tags.extend(re.findall(jinja_pattern, shape.text))
    
    return jinja_tags

# Пример использования
tags = extract_jinja_tags('examples/3-ПТ2-1.docx')
for tag in tags:
    print(tag)