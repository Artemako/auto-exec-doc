import re
from docx import Document

def find_jinja_tags(docx_path):
    # Загружаем документ
    doc = Document(docx_path)
    jinja_tags = []
    # Регулярное выражение для поиска тегов Jinja
    jinja_pattern = re.compile(r'\{\{.*?\}\}|\{%.*?%\}')

    # Функция для поиска тегов в текстах
    def search_tags_in_texts(texts):
        for text in texts:
            matches = jinja_pattern.findall(text)
            jinja_tags.extend(matches)

    # Ищем теги в параграфах
    search_tags_in_texts(paragraph.text for paragraph in doc.paragraphs)
    
    # Ищем теги в колонтитулах
    for section in doc.sections:
        search_tags_in_texts(paragraph.text for paragraph in section.header.paragraphs)  # В верхнем колонтитуле
        search_tags_in_texts(paragraph.text for paragraph in section.footer.paragraphs)  # В нижнем колонтитуле

    # Ищем теги в фигурных элементах документа
    for shape in doc.inline_shapes:
        if shape.type == 1:  # Проверка на тип "текстовое поле" (1 соответствует текстовым полям)
            search_tags_in_texts([shape.text])

    return set(jinja_tags)

# Пример использования
docx_file = 'C:/Users/hayar/Documents/AutoExecDoc Projects/gdgffdg/forms/10.docx'
tags = find_jinja_tags(docx_file)
print("Найденные Jinja теги:")
for tag in tags:
    print(tag)
