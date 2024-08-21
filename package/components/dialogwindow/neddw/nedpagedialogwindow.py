from PySide6.QtWidgets import (
    QDialog,
    QTableWidget,
    QTableWidgetItem,
    QAbstractItemView,
    QPushButton,
    QHBoxLayout,
    QWidget,
    QHeaderView,
)

import package.ui.nedpagedialogwindow_ui as nedpagedialogwindow_ui

import package.components.dialogwindow.neddw.nedvariabledialogwindow as nedvariabledialogwindow

import json
import os
import datetime
import re
from docx import Document
from functools import partial


class NedPageDialogWindow(QDialog):
    def __init__(self, osbm, type_ned, pages, page=None):
        self.__osbm = osbm
        self.__type_ned = type_ned
        self.__pages = pages
        self.__page = page
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow __init__(osbm, type_ned, page):\ntype_ned = {self.__type_ned}\npages = {self.__pages}\npage = {self.__page}"
        )
        super(NedPageDialogWindow, self).__init__()
        self.ui = nedpagedialogwindow_ui.Ui_NedPageDialogWindow()
        self.ui.setupUi(self)
        # СТИЛЬ
        self.__osbm.obj_style.set_style_for(self)
        #
        self.__page_filename = str()
        self.__temp_file_path = str()
        self.__data = {
            "id_parent_template": None,
            "name_page": None,
            "filename_page": None,
            "order_page": None,
            "included": 1,
        }
        self.__variables_for_add = []
        # одноразовые действия
        self.config_by_type_window()
        self.config_combox_neighboor()
        self.reconfig_tw_variables()
        self.connecting_actions()

    def get_data(self):
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow get_data():\nself.__data = {self.__data}"
        )
        return self.__data

    def config_by_type_window(self):
        self.__osbm.obj_logg.debug_logger("NedPageDialogWindow config_by_type_window()")
        if self.__type_ned == "create":
            self.ui.btn_select.setText("Выбрать документ")
            self.ui.btn_open_docx.setEnabled(False)
            self.ui.label_file.setText("Файл не выбран")
            self.ui.btn_nesvariable.setText("Добавить страницу")
        elif self.__type_ned == "edit":
            self.ui.btn_select.setText("Выбрать новый документ")
            self.ui.btn_open_docx.setEnabled(True)
            self.ui.label_file.setText(self.__page.get("filename_page"))
            self.__page_filename = self.__page.get("filename_page")
            self.ui.btn_nesvariable.setText("Сохранить страницу")
            self.ui.lineedit_namepage.setText(self.__page.get("name_page"))

    def reconfig_tw_variables(self):
        self.__osbm.obj_logg.debug_logger("NedPageDialogWindow reconfig_tw_variables()")
        tablewidget = self.ui.tw_variables
        tablewidget.blockSignals(True)
        tablewidget.clearContents()
        tablewidget.setRowCount(0)
        if self.__temp_file_path:
            jinja_variables = self.extract_jinja_variables(self.__temp_file_path)
            if jinja_variables:
                self.fill_tw_variables(jinja_variables)
        elif self.__page_filename:
            forms_folder_dirpath = self.__osbm.obj_dirm.get_forms_folder_dirpath()
            docx_path = os.path.join(
                forms_folder_dirpath, self.__page_filename + ".docx"
            )
            jinja_variables = self.extract_jinja_variables(docx_path)
            if jinja_variables:
                self.fill_tw_variables(jinja_variables)
        tablewidget.blockSignals(False)

    def fill_tw_variables(self, jinja_variables):
        # TODO ПЕРЕДЕЛАТЬ
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow fill_tw_variables(jinja_variables):\njinja_variables = {jinja_variables}"
        )
        tablewidget = self.ui.tw_variables
        tablewidget.setColumnCount(3)
        tablewidget.setHorizontalHeaderLabels(["Переменная", "Вид переменной", "Действия"])
        tablewidget.setRowCount(len(jinja_variables))
        for row, variable in enumerate(jinja_variables):
            qtwt_name_variable = QTableWidgetItem(variable)
            # qtwt_name_variable.setData(1001, variable)
            status_variable = self.get_status_variable(variable)
            type_variable = status_variable.get("type")
            value_variable = status_variable.get("value")
            qtwt_status_variable = QTableWidgetItem(type_variable)
            # кнопка
            layout = QHBoxLayout()
            add_button = QPushButton("...")
            layout.setContentsMargins(0, 0, 0, 0)
            layout.addWidget(add_button)
            widget = QWidget()
            widget.setLayout(layout)
            #
            tablewidget.setItem(row, 0, qtwt_name_variable)
            tablewidget.setItem(row, 1, qtwt_status_variable)
            #
            if type_variable == "Переменная":
                result_bd = self.__osbm.obj_prodb.get_variable_by_name(value_variable)
                result_variables_for_add = self.get_variable_by_name(value_variable)
                if result_bd or result_variables_for_add:
                    tablewidget.setItem(row, 2, QTableWidgetItem("Имеется"))
                else:
                    add_button.setText("Добавить переменную")
                    add_button.clicked.connect(partial(self.add_variable, value_variable, False))
                    tablewidget.setCellWidget(row, 2, widget)
            elif type_variable == "Блок" and value_variable:
                result_bd = self.__osbm.obj_prodb.get_variable_by_name(value_variable)
                result_variables_for_add = self.get_variable_by_name(value_variable)
                if result_bd or result_variables_for_add:
                    tablewidget.setItem(row, 2, QTableWidgetItem("Имеется"))
                else:
                    add_button.setText("Добавить блок")
                    add_button.clicked.connect(
                        partial(self.add_variable, value_variable, True)
                    )
                    tablewidget.setCellWidget(row, 2, widget)
        # Настраиваем режимы изменения размера для заголовков
        header = tablewidget.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        #
        tablewidget.setSortingEnabled(True)
        tablewidget.resizeColumnsToContents()
        tablewidget.setEditTriggers(QTableWidget.NoEditTriggers)
        tablewidget.setSelectionMode(QAbstractItemView.NoSelection)

    def get_variable_by_name(self, name_variable) -> dict:
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow get_variable_by_name(name_variable) -> dict:\nname_variable = {name_variable}"
        )
        variables = self.__variables_for_add
        result = None
        for variable in variables:
            if variable.get("name_variable") == name_variable:
                result = variable
                break
        return result

    def get_status_variable(self, variable: str) -> dict:
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow get_status_variable(variable) -> dict:\nvariable = {variable}"
        )
        status_variable = {"type": None, "value": None}
        # statuses = ["Тэг", "Атрибут тега", "Блок", "Прочее"]
        if variable.startswith("{%") and variable.endswith("%}"):
            status_variable["type"] = "Блок"
            words_in_variable = variable.split()
            if words_in_variable[-2] != "endfor":
                status_variable["value"] = words_in_variable[-2]
        elif variable.startswith("{{") and variable.endswith("}}"):
            if "." in variable:
                status_variable["type"] = "Атрибут тега"
            else:
                status_variable["type"] = "Переменная"
                words_in_variable = variable.split()
                status_variable["value"] = words_in_variable[1]
        else:
            status_variable["type"] = "Прочее"
        return status_variable

    def add_variable(self, name_variable, is_block):
        """
        Похожий код у VariablesListDialogWindow.
        """
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow add_variable(variable, name_variable, is_block):\nname_variable = {name_variable}\n is_block = {is_block}"
        )
        #
        self.__all_variables = self.__osbm.obj_prodb.get_variables()
        # окно        
        self.__osbm.obj_nedtdw = nedvariabledialogwindow.NedVariableDialogWindow(
            self.__osbm, "create", self.__all_variables, None, name_variable, is_block
        )
        result = self.__osbm.obj_nedtdw.exec()
        # результат
        if result == QDialog.Accepted:
            # получить data
            data = self.__osbm.obj_nedtdw.get_data()
            name_variable = data.get("NAME")
            type_variable = data.get("TYPE")
            title_variable = data.get("TITLE")
            order_variable = data.get("ORDER")
            config_variable = data.get("CONFIG")
            config_variable = json.dumps(config_variable) if config_variable else None
            description_variable = data.get("DESCRIPTION")
            create_variable = {
                "name_variable": name_variable,
                "type_variable": type_variable,
                "title_variable": title_variable,
                "order_variable": order_variable,
                "config_variable": config_variable,
                "description_variable": description_variable,
            }
            # вставка
            self.__all_variables.insert(order_variable, create_variable)
            self.__osbm.obj_prodb.insert_variable(create_variable)
            # Меняем порядок в БД - кому нужно
            for index, variable in enumerate(self.__all_variables):
                order = variable.get("order_variable")
                if order != index:
                    self.__osbm.obj_prodb.set_order_for_variable(variable, index)
            # обновляем таблицу
            self.reconfig_tw_variables()



    def extract_jinja_variables(self, docx_path) -> set:
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow extract_jinja_variables(docx_path) -> set:\ndocx_path = {docx_path}"
        )
        try:
            doc = Document(docx_path)
            # Регулярное выражение для поиска всех Jinja переменных
            # TODO Убрать пробелы + найти именно все-все теги
            jinja_pattern = r"\{\%\s.*?\%\}|\{\{\s.*?\s\}\}"
            jinja_variables = []
            # Обходим все параграфы в документе
            for para in doc.paragraphs:
                jinja_variables.extend(re.findall(jinja_pattern, para.text))
            # Обходим все таблицы в документе
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        jinja_variables.extend(re.findall(jinja_pattern, cell.text))
            # Обходим колонтитулы
            for section in doc.sections:
                for header in section.header.paragraphs:
                    jinja_variables.extend(re.findall(jinja_pattern, header.text))
                for footer in section.footer.paragraphs:
                    jinja_variables.extend(re.findall(jinja_pattern, footer.text))
            # Поиск в фигурах (шапках с рисованием, если поддерживаются)
            for shape in doc.inline_shapes:
                if shape.type == 3:  # 3 - тип 'Picture'
                    continue
                jinja_variables.extend(re.findall(jinja_pattern, shape.text))
            return set(jinja_variables)
        except Exception as e:
            self.__osbm.obj_logg.error_logger(f"Error in extract_jinja_variables: {e}")
            return set()

    def config_combox_neighboor(self):
        self.__osbm.obj_logg.debug_logger(
            "NedPageDialogWindow config_combox_neighboor()"
        )
        combobox = self.ui.combox_neighboor
        combobox.blockSignals(True)
        combobox.clear()
        current_index = 0
        # по умолчанию - в конец
        flag = True
        combobox.addItem("- В начало -", "START")
        for index, page in enumerate(self.__pages):
            if self.__page and self.__page.get("id_page") == page.get("id_page"):
                flag = False
            else:
                combobox.addItem(page.get("name_page"), page)
            if flag:
                current_index = index + 1
        combobox.setCurrentIndex(current_index)
        combobox.blockSignals(False)

    def connecting_actions(self):
        self.__osbm.obj_logg.debug_logger("NedPageDialogWindow connecting_actions()")
        self.ui.btn_select.clicked.connect(self.select_file)
        self.ui.btn_open_docx.clicked.connect(self.open_docx)
        self.ui.btn_nesvariable.clicked.connect(self.btn_nesvariable_clicked)
        self.ui.btn_nesvariable.setShortcut("Ctrl+S")
        self.ui.btn_close.clicked.connect(self.close)
        self.ui.btn_close.setShortcut("Ctrl+Q")

    def select_file(self):
        docx_path = self.__osbm.obj_dw.select_docx_file()
        if docx_path:
            # текст
            self.ui.label_file.setText(os.path.basename(docx_path))
            # образуем название документа
            file_name = f"docx_{datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S')}"
            self.__page_filename = file_name
            file_name_with_docx = f"{file_name}.docx"
            # путь к временной папке
            temp_dir = self.__osbm.obj_dirm.get_temp_dirpath()
            # путь к временному файлу
            self.__temp_file_path = os.path.join(temp_dir, file_name_with_docx)
            # копирование
            self.__osbm.obj_film.copy_file(docx_path, self.__temp_file_path)
            #
            self.reconfig_tw_variables()

    def open_docx(self):
        # todo ПРИ CREATE заблокировать кнопку Открыть и изменить документ
        # TODO После закрытия формы обновить теги (кнопку обновить/поиск добавить)
        self.__osbm.obj_logg.debug_logger("NedPageDialogWindow open_docx()")
        try:
            # название docx
            filename_page = self.__page.get("filename_page")
            # путь к документу
            forms_folder_dirpath = self.__osbm.obj_dirm.get_forms_folder_dirpath()
            docx_path = os.path.join(forms_folder_dirpath, filename_page + ".docx")
            print(f"docx_path = {docx_path}")
            # открытие
            if os.path.exists(docx_path):
                try:
                    os.startfile(docx_path)
                except OSError:
                    raise Exception("Не удалось открыть.")
        except Exception as e:
            self.__osbm.obj_dw.warning_message(f"Error: {e}")
            self.__osbm.obj_dw.warning_message("Открыть не удалось.")
            print(e)

    def find_page_by_namepage_in_pages(self, namepage) -> object:
        self.__osbm.obj_logg.debug_logger(
            f"NedPageDialogWindow find_page_by_namepage_in_pages(namepage):\nnamepage = {namepage}"
        )
        for current_page in self.__pages:
            if current_page.get("name_page") == namepage:
                return current_page
        return None

    def btn_nesvariable_clicked(self):
        self.__osbm.obj_logg.debug_logger("NedPageDialogWindow btn_nesvariable_clicked()")

        filename_page = self.__page_filename
        namepage = self.ui.lineedit_namepage.text()
        if len(namepage) > 0 and len(filename_page) > 0:
            # поиск по имени
            find_page = self.find_page_by_namepage_in_pages(namepage)
            # выбранный сосед
            neighboor_page = self.ui.combox_neighboor.currentData()
            order_page = (
                0 if neighboor_page == "START" else neighboor_page.get("order_page") + 1
            )
            #
            if self.__type_ned == "create":
                if find_page is None:
                    self.__data["name_page"] = namepage
                    self.__data["filename_page"] = filename_page
                    self.__data["order_page"] = order_page
                    self.accept()
                else:
                    msg = "Другая страница с таким именем уже существует!"
                    self.__osbm.obj_dw.warning_message(msg)

            elif self.__type_ned == "edit":
                if find_page is None:
                    self.__data = self.__page
                    self.__data["name_page"] = namepage
                    self.__data["filename_page"] = filename_page
                    self.__data["order_page"] = order_page
                    self.accept()
                elif namepage == self.__page.get("name_page"):
                    self.__data = self.__page
                    self.__data["name_page"] = namepage
                    self.__data["filename_page"] = filename_page
                    self.__data["order_page"] = order_page
                    self.accept()
                else:
                    msg = "Другая страница с таким именем уже существует!"
                    self.__osbm.obj_dw.warning_message(msg)
        elif namepage == "":
            self.__osbm.obj_dw.warning_message("Заполните поле названия")
        elif filename_page is None or len(filename_page) == 0:
            self.__osbm.obj_dw.warning_message("Выберите документ")
        else:
            self.__osbm.obj_dw.warning_message(
                "Заполните поле названия и выберите документ"
            )
