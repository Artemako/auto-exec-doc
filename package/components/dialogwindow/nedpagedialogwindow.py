from PySide6.QtWidgets import (
    QDialog,
    QTableWidget,
    QTableWidgetItem,
    QAbstractItemView,
    QPushButton,
    QSizePolicy,
    QHBoxLayout,
    QWidget,
    QHeaderView
)
from PySide6.QtCore import QTimer, QSize
from PySide6.QtGui import QIcon

import package.ui.nedpagedialogwindow_ui as nedpagedialogwindow_ui

import os
from pathlib import Path
import datetime
import re
from docx import Document
from functools import partial


class NedPageDialogWindow(QDialog):
    def __init__(self, obs_manager, type_ned, pages, page=None):
        self.__obs_manager = obs_manager
        self.__type_ned = type_ned
        self.__pages = pages
        self.__page = page
        self.__obs_manager.obj_l.debug_logger(
            f"NedPageDialogWindow __init__(obs_manager, type_ned, page):\ntype_ned = {self.__type_ned}\npages = {self.__pages}\npage = {self.__page}"
        )
        super(NedPageDialogWindow, self).__init__()
        self.ui = nedpagedialogwindow_ui.Ui_NedPageDialogWindow()
        self.ui.setupUi(self)
        # СТИЛЬ
        self.__obs_manager.obj_style.set_style_for(self)
        #
        self.__page_filename = str()
        self.__data = {
            "id_parent_template": None,
            "name_page": None,
            "filename_page": None,
            "order_page": None,
            "included": 1,
        }
        self.__tags_for_add = []
        self.__icons = self.__obs_manager.obj_icons.get_icons()
        # одноразовые действия
        self.config_by_type_window()
        self.config_combox_neighboor()
        self.reconfig_tw_tags()
        self.connecting_actions()

    def config_by_type_window(self):
        self.__obs_manager.obj_l.debug_logger(
            "NedPageDialogWindow config_by_type_window()"
        )
        if self.__type_ned == "create":
            self.ui.btn_select.setText("Выбрать документ")
            self.ui.btn_open_docx.setEnabled(False)
            self.ui.label_file.setText("Файл не выбран")
            self.ui.btn_nestag.setText("Добавить страницу")
            self.ui.btn_nestag.setIcon(self.__icons.get("add"))
        elif self.__type_ned == "edit":
            self.ui.btn_select.setText("Выбрать новый документ")
            self.ui.btn_open_docx.setEnabled(True)
            self.ui.label_file.setText(self.__page.get("filename_page"))
            self.__page_filename = self.__page.get("filename_page")
            self.ui.btn_nestag.setText("Сохранить страницу")
            self.ui.btn_nestag.setIcon(self.__icons.get("save"))
            self.ui.lineedit_namepage.setText(self.__page.get("name_page"))


    def reconfig_tw_tags(self):
        self.__obs_manager.obj_l.debug_logger("NedPageDialogWindow reconfig_tw_tags()")
        if self.__type_ned == "edit":
            tablewidget = self.ui.tw_tags
            tablewidget.blockSignals(True)
            tablewidget.clearContents()
            tablewidget.setRowCount(0)
            if self.__page_filename:
                forms_folder_dirpath = self.__obs_manager.obj_dpm.get_forms_folder_dirpath()
                docx_path = os.path.join(
                    forms_folder_dirpath, self.__page_filename + ".docx"
                )
                jinja_tags = self.extract_jinja_tags(docx_path)
                if jinja_tags:
                    self.fill_tw_tags(jinja_tags)
            tablewidget.blockSignals(False)

    def fill_tw_tags(self, jinja_tags):
        self.__obs_manager.obj_l.debug_logger(f"NedPageDialogWindow fill_tw_tags(jinja_tags):\njinja_tags = {jinja_tags}")
        tablewidget = self.ui.tw_tags
        tablewidget.setColumnCount(3)
        tablewidget.setHorizontalHeaderLabels(["Тег", "Вид тега", "Действия"])
        tablewidget.setRowCount(len(jinja_tags))
        for row, tag in enumerate(jinja_tags):
            qtwt_name_tag = QTableWidgetItem(tag)
            # qtwt_name_tag.setData(1001, tag)
            status_tag = self.get_status_tag(tag)
            type_tag = status_tag.get("type")
            value_tag = status_tag.get("value")
            qtwt_status_tag = QTableWidgetItem(type_tag)
            # кнопка 
            layout = QHBoxLayout()
            add_button = QPushButton("...")
            layout.setContentsMargins(0, 0, 0, 0)
            layout.addWidget(add_button)
            widget = QWidget()
            widget.setLayout(layout)
            #
            tablewidget.setItem(row, 0, qtwt_name_tag)
            tablewidget.setItem(row, 1, qtwt_status_tag)
            #
            if type_tag == "Переменная":
                result_bd = self.__obs_manager.obj_pd.get_tag_by_name(value_tag)
                result_tags_for_add = self.get_tag_by_name(value_tag)
                if result_bd or result_tags_for_add:
                    tablewidget.setItem(row, 2, QTableWidgetItem("Имеется"))
                else:
                    add_button.setText("Добавить тег")
                    add_button.clicked.connect(partial(self.add_tag, tag, value_tag))
                    tablewidget.setCellWidget(row, 2, widget)
            elif type_tag == "Блок" and value_tag:
                result_bd = self.__obs_manager.obj_pd.get_tag_by_name(value_tag)
                result_tags_for_add = self.get_tag_by_name(value_tag)
                if result_bd or result_tags_for_add:
                    tablewidget.setItem(row, 2, QTableWidgetItem("Имеется"))
                else:
                    add_button.setText("Добавить блок")
                    add_button.clicked.connect(partial(self.add_tag, tag, value_tag, is_block = True))
                    tablewidget.setCellWidget(row, 2, widget)
        # Настраиваем режимы изменения размера для заголовков
        header = tablewidget.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Stretch)
        header.setSectionResizeMode(1, QHeaderView.ResizeToContents)        
        header.setSectionResizeMode(2, QHeaderView.ResizeToContents)
        #
        tablewidget.setSortingEnabled(True)        
        tablewidget.resizeColumnsToContents()
        tablewidget.setEditTriggers(QTableWidget.NoEditTriggers)
        tablewidget.setSelectionMode(QAbstractItemView.NoSelection) 

    def get_tag_by_name(self, name_tag) -> dict:
        self.__obs_manager.obj_l.debug_logger(f"NedPageDialogWindow get_tag_by_name(name_tag) -> dict:\nname_tag = {name_tag}")
        tags = self.__tags_for_add
        result = None
        for tag in tags:
            if tag.get("name_tag") == name_tag:
                result = tag
                break
        return result

    def get_status_tag(self, tag : str) -> dict:
        self.__obs_manager.obj_l.debug_logger(f"NedPageDialogWindow get_status_tag(tag) -> dict:\ntag = {tag}")
        status_tag = {
            "type" : None,
            "value" : None
        }
        # statuses = ["Тэг", "Атрибут тега", "Блок", "Прочее"]
        if tag.startswith("{%") and tag.endswith("%}"):
            status_tag["type"] = "Блок"
            words_in_tag = tag.split()
            if words_in_tag[-2] != "endfor":
                status_tag["value"] = words_in_tag[-2]
        elif tag.startswith("{{") and tag.endswith("}}"):
            if "." in tag:
                status_tag["type"] = "Атрибут тега"
            else:
                status_tag["type"] = "Переменная"
                words_in_tag = tag.split()
                status_tag["value"] = words_in_tag[1]
        else:
            status_tag["type"] = "Прочее"
        return status_tag


    def add_tag(self, tag, name_tag = None, is_block = False):
        self.__obs_manager.obj_l.debug_logger(f"NedPageDialogWindow add_tag(tag, name_tag, is_block):\ntag = {tag}\nname_tag = {name_tag}\n is_block = {is_block}")
        # TODO

    def extract_jinja_tags(self, docx_path) -> set:
        self.__obs_manager.obj_l.debug_logger(f"NedPageDialogWindow extract_jinja_tags(docx_path) -> set:\ndocx_path = {docx_path}")
        try:
            doc = Document(docx_path)
            # Регулярное выражение для поиска Jinja тегов
            jinja_pattern = r"\{\%.*?\%\}|\{\{.*?\}\}"
            jinja_tags = []
            # Обходим все параграфы в документе
            for para in doc.paragraphs:
                jinja_tags.extend(re.findall(jinja_pattern, para.text))
            # Обходим все таблицы в документе
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        jinja_tags.extend(re.findall(jinja_pattern, cell.text))
            # Обходим колонтитулы
            for section in doc.sections:
                for header in section.header.paragraphs:
                    jinja_tags.extend(re.findall(jinja_pattern, header.text))
                for footer in section.footer.paragraphs:
                    jinja_tags.extend(re.findall(jinja_pattern, footer.text))
            # Поиск в фигурах (шапках с рисованием, если поддерживаются)
            for shape in doc.inline_shapes:
                if shape.type == 3:  # 3 - тип 'Picture'
                    continue
                jinja_tags.extend(re.findall(jinja_pattern, shape.text))
            return set(jinja_tags)
        except Exception as e:
            self.__obs_manager.obj_l.error_logger(f"Error in extract_jinja_tags: {e}")
            return set()

    def config_combox_neighboor(self):
        self.__obs_manager.obj_l.debug_logger(
            "NedPageDialogWindow config_combox_neighboor()"
        )
        combobox = self.ui.combox_neighboor
        combobox.blockSignals(True)
        combobox.clear()
        current_index = 0
        flag = True
        combobox.addItem("- В начало -", "start")
        for index, page in enumerate(self.__pages):
            if self.__page.get("id_page") == page.get("id_page"):
                flag = False
            else:
                combobox.addItem(page.get("name_page"), page)
            if flag:
                current_index = index + 1
        combobox.setCurrentIndex(current_index)
        combobox.blockSignals(False)

    def connecting_actions(self):
        self.__obs_manager.obj_l.debug_logger(
            "NedPageDialogWindow connecting_actions()"
        )
        self.ui.btn_select.clicked.connect(self.select_file)
        self.ui.btn_open_docx.clicked.connect(self.open_docx)
        self.ui.btn_nestag.clicked.connect(self.btn_nestag_clicked)
        self.ui.btn_nestag.setShortcut("Ctrl+S")
        self.ui.btn_close.clicked.connect(self.close)
        self.ui.btn_close.setShortcut("Ctrl+Q")

    def select_file(self):
        docx_path = self.__obs_manager.obj_dw.select_docx_file()
        if docx_path:
            # текст
            self.ui.label_file.setText(os.path.basename(docx_path))
            # образуем название документа
            file_name = f"docx_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}"
            self.__page_filename = file_name
            file_name_with_docx = f"{file_name}.docx"
            # путь к временной папке
            temp_dir = self.__obs_manager.obj_dpm.get_temp_dirpath()
            # путь к временному файлу
            temp_file_path = os.path.join(temp_dir, file_name_with_docx)
            # копирование
            self.__obs_manager.obj_dpm.copy_file(docx_path, temp_file_path)
            #
            self.reconfig_tw_tags()

    def open_docx(self):
        self.__obs_manager.obj_l.debug_logger("NedPageDialogWindow open_docx()")
        try:
            # название docx
            filename_page = self.__page.get("filename_page") 
            # путь к документу
            forms_folder_dirpath = self.__obs_manager.obj_dpm.get_forms_folder_dirpath()
            docx_path = os.path.join(
                forms_folder_dirpath, filename_page + ".docx"
            )
            # открытие
            if os.path.exists(docx_path):
                try:
                    os.startfile(docx_path)
                except OSError:
                    raise Exception("Не удалось открыть.")
        except Exception as e:
            self.__obs_manager.obj_dw.warning_message(f"Error: {e}")
            self.__obs_manager.obj_dw.warning_message("Открыть не удалось.")
            print(e)


    def btn_nestag_clicked(self):
        self.__obs_manager.obj_l.debug_logger(
            "NedPageDialogWindow btn_nestag_clicked()"
        )
        if self.__type_ned == "create":
            filename_page = self.__page_filename
            namepage = self.ui.lineedit_namepage.text()
            # TODO Уникальность имени!!!
            if len(namepage) > 0 and len(filename_page) > 0:
                self.__data["name_page"] = namepage
                self.__data["filename_page"] = filename_page
                self.accept()
            elif namepage == "":
                self.__obs_manager.obj_dw.warning_message("Заполните поле названия")
            elif filename_page is None or len(filename_page) == 0:
                self.__obs_manager.obj_dw.warning_message("Выберите документ")
            else:
                self.__obs_manager.obj_dw.warning_message(
                    "Заполните поле названия и выберите документ"
                )
        elif self.__type_ned == "edit":
            self.__data = self.__page
            # TODO edit
